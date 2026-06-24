"""Generate desktop Word user manual from docs/USAGE.md."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
USAGE_MD = ROOT / "docs" / "USAGE.md"
DESKTOP_DOCX = Path.home() / "Desktop" / "亚马逊广告诊断助手_使用说明书.docx"

APP_URL = "https://amazon-ad-assistant-2.streamlit.app"
GITHUB_URL = "https://github.com/John-Z-jt/amazon-ad-assistant-2"


def add_hyperlink_paragraph(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}{url}")
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run.font.size = Pt(11)


def parse_table_lines(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    header = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        if not line.strip().startswith("|"):
            break
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return header, rows


def build_doc() -> Document:
    text = USAGE_MD.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1
            header, rows = parse_table_lines(table_lines)
            if header:
                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = "Table Grid"
                for col, val in enumerate(header):
                    table.rows[0].cells[col].text = val
                for r, row in enumerate(rows):
                    for col, val in enumerate(row):
                        table.rows[r + 1].cells[col].text = val
                doc.add_paragraph()
            i = j
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph("\n".join(code_lines))
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:].strip())
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        if stripped.startswith("**在线使用地址：**"):
            doc.add_paragraph()
            add_hyperlink_paragraph(doc, "在线使用地址：", APP_URL)
            i += 1
            continue

        if stripped.startswith("**开源仓库"):
            add_hyperlink_paragraph(doc, "开源仓库：", GITHUB_URL)
            i += 1
            continue

        # bold-only lines like **Q：...**
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        clean = clean.replace("`", "")
        doc.add_paragraph(clean)
        i += 1

    return doc


if __name__ == "__main__":
    if not USAGE_MD.exists():
        raise SystemExit(f"Missing {USAGE_MD}")
    doc = build_doc()
    doc.save(DESKTOP_DOCX)
    print(f"Saved: {DESKTOP_DOCX}")
